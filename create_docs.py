# Load modules
import csv
import json
from docx import Document
from docx.shared import RGBColor
import re

EM_DASH = u'\u2014'

# Clean up trailing zeros
def trailing_zeros(val):
    '''Clean up trailing zeroes (e.g. converts 4.00 to 4)'''
    val = val.replace('.00', '')
    val = re.sub('(.[1-9])0', r'\1', val)
    return val

# Get year from join date
def get_join_yr(join_date):
    '''Get the year from the join date string'''
    date_arr = join_date.split('/')
    if len(date_arr) > 1:
        return date_arr[2]
    return join_date

# Reset Error Log
with open("errorlog.txt", "w") as f:
    f.write("")
f.close()

# Read data from JSON file
file_info = json.load(open('info.json'))
import_addr = file_info['dir'] + '\\' + file_info['filename']

# Creates a list of dictionaries from the CSV data,
# excluding any that are marked as having the Word doc done
with open(import_addr, newline='') as csv_file:
    import_data = [
        row for row in csv.DictReader(csv_file)
        if row['Word Doc Done'] == 'No'
    ]
csv_file.close()

# Sets up the fields that will turn into a bulleted list
bullet_fields = ['Entry Feature 1', 'Entry Feature 2', 'Entry Feature 3', 'Entry Feature 4', 'Entry Feature 5', 'Entry Feature 6']

# Sets up the fields that will turn into the table
table_data = [
    {'label' : 'Finished SqFt', 'field' : 'Living SF'},
    {'label' : 'Garage SqFt', 'field' : 'Garage SF'},
    {'label' : 'Total SqFt', 'field' : 'Total SF'},
    {'label' : 'Number of Bedrooms', 'field' : 'Bedrooms'},
    {'label' : 'Number of Baths', 'field' : 'Bathrooms'}
]

# Writes Word documents
for entry in import_data:

    # Create a Word doc
    document = Document()

    # Write entry number placeholder
    paragraph = document.add_paragraph()
    paragraph.add_run('#').bold = True
    run = paragraph.add_run('xx')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True

    # Write price
    paragraph = document.add_paragraph()
    paragraph.add_run(entry['Price'])

    # Write company info block
    paragraph = document.add_paragraph()
    paragraph.style = 'No Spacing'
    paragraph.add_run(entry['Attendee Company']).bold = True
    paragraph = document.add_paragraph()
    paragraph.style = 'No Spacing'
    paragraph.add_run(entry['Phone'])
    paragraph = document.add_paragraph()
    paragraph.add_run(entry['Website'])

    # Write the address
    paragraph = document.add_paragraph()
    paragraph.add_run(entry['Entry Street Address'] + ' {0} '.format(EM_DASH) + entry['Entry City']).bold = True

    # Write subdivision
    paragraph = document.add_paragraph()
    paragraph.add_run('Subdivision: ').bold = True
    paragraph.add_run(entry['Entry Subdivision'])

    # Write description
    paragraph = document.add_paragraph()
    paragraph.add_run('Description: ').bold = True
    paragraph.add_run(entry['Description'])

    # Add bulleted list of features
    for field in bullet_fields:
        if len(entry[field]) > 0:
            document.add_paragraph(entry[field], style='List Bullet')

    if entry['Entry Type'] == 'Single-Family Home' or entry['Entry Type'] == 'Townhome Unit':
        # Write home style
        paragraph = document.add_paragraph()
        paragraph.add_run('Home Style: ').bold = True
        paragraph.add_run(entry['Home Style'])
    
        # Write exterior
        paragraph = document.add_paragraph()
        paragraph.add_run('Exterior: ').bold = True
        paragraph.add_run(entry['Exterior'])

        # Write HERS
        paragraph = document.add_paragraph()
        paragraph.add_run('HERS: ').bold = True
        if len(entry['HERS Rating']) != 0:
            paragraph.add_run(entry['HERS Rating'])
        else:
            run = paragraph.add_run('Not Yet Assessed')
            run.font.color.rgb = RGBColor(255, 0, 0)

        # Create table
        table = document.add_table(rows=5, cols=2)
        for i, data in enumerate(table_data):
            row = table.rows[i].cells
            row[0].text = data['label']
            table_value = entry[data['field']]
            if i > 2:
                table_value = trailing_zeros(table_value)
            row[1].text = table_value

        # Write contractor license number
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph()
        paragraph.add_run('Contractor License #' + entry['Contractor License']).italic = True

    else:
        # Write entry type
        paragraph = document.add_paragraph()
        paragraph.add_run('Entry Type: ').bold = True
        paragraph.add_run(entry['Entry Type'])

    # Write member since info
    paragraph = document.add_paragraph()
    join_yr = get_join_yr(entry['Member Join Date'])
    paragraph.add_run('This company has been a Member since ' + join_yr + '.').italic = True

    # Save the document
    document.save(file_info['dir'] + '\\' + entry['Entry Street Address'] + ' - Magazine Info.docx')