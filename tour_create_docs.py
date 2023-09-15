# Load modules
import csv
import json
from docx import Document
from docx.shared import RGBColor
import re

EM_DASH = u'\u2014'

# Clean up trailing zeros
def trailing_zeros(val):
    '''Clean up trailing zeroes (e.g. converts 4.00 to 4)'''
    val = val.replace('.00', '')
    val = re.sub('(.[1-9])0', r'\1', val)
    return val

# Get year from join date
def get_join_yr(join_date):
    '''Get the year from the join date string'''
    date_arr = join_date.split('/')
    if len(date_arr) > 1:
        return date_arr[2]
    return join_date

# Reset Error Log
with open("errorlog.txt", "w") as f:
    f.write("")
f.close()

# Read data from JSON file
file_info = json.load(open('info.json'))['tour']
import_addr = file_info['dir'] + '\\' + file_info['filename']
exclude_addr = file_info['dir'] + '\\' + file_info['exclude']

# Gets a list of registration IDs to exclude
with open(exclude_addr, newline='') as csv_file:
    exclude_data = set(
        row['RegID'] for row in csv.DictReader(csv_file)
    )
csv_file.close()

# Creates a list of dictionaries from the CSV data,
# excluding any that are marked as having the Word doc done
with open(import_addr, newline='') as csv_file:
    import_data = [
        row for row in csv.DictReader(csv_file)
        if row['Word Doc Done'] == 'No'
        and row['RegID'] not in exclude_data
    ]
csv_file.close()

# Sets up the fields that will turn into a bulleted list
bullet_fields = ['Project Goal 1', 'Project Goal 2', 'Project Goal 3', 'Project Goal 4']

# Writes Word documents
for entry in import_data:

    # Create a Word doc
    document = Document()

    # Write entry number placeholder
    paragraph = document.add_paragraph()
    paragraph.add_run('#').bold = True
    run = paragraph.add_run('xx')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True

    # Write price
    paragraph = document.add_paragraph()
    paragraph.add_run(entry['Remodel Type'])

    # Write company info block
    paragraph = document.add_paragraph()
    paragraph.style = 'No Spacing'
    paragraph.add_run(entry['Attendee Company']).bold = True
    paragraph = document.add_paragraph()
    paragraph.style = 'No Spacing'
    paragraph.add_run(entry['Phone'])
    paragraph = document.add_paragraph()
    paragraph.add_run(entry['Website'])

    # Write the address
    paragraph = document.add_paragraph()
    # paragraph.add_run(entry['Entry Street Address'] + ' {0} '.format(EM_DASH) + entry['Entry City']).bold = True
    paragraph.add_run(entry['Entry Street Address'] + ' ' + entry['Entry City']).bold = True

    # Write description
    paragraph = document.add_paragraph()
    paragraph.add_run('Description: ').bold = True
    paragraph.add_run(entry['Description'])

    # Add bulleted list of project goals
    for field in bullet_fields:
        if len(entry[field]) > 0:
            document.add_paragraph(entry[field], style='No Spacing')
    # Add extra space (for non-bullet style)
    document.add_paragraph('', style='No Spacing')
            
    # Write contractor license number
    paragraph = document.add_paragraph()
    paragraph = document.add_paragraph()
    # Add support for N/A (since contractors may not have license)
    if len(entry['Contractor License']) < 2:
        entry['Contractor License'] = "N/A"
    paragraph.add_run('Contractor License #' + entry['Contractor License']).italic = True

    # Write member since info
    paragraph = document.add_paragraph()
    join_yr = get_join_yr(entry['Member Join Date'])
    paragraph.add_run('This company has been a Member since ' + join_yr + '.').italic = True

    # Save the document
    document.save(file_info['dir'] + '\\' + entry['Entry Street Address'] + ' - Magazine Info.docx')