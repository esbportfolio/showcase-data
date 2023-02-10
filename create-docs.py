# Load modules
import csv
from dotenv import load_dotenv
import os
from docx import Document
from docx.shared import RGBColor
import re

em_dash = u'\u2014'

# Clean up trailing zeros
def trailing_zeros(val):
    val = val.replace('.00', '')
    val = re.sub('(.[1-9])0', r'\1', val)
    return val

# Get year from join date
def get_join_yr(join_date):
    date_arr = join_date.split('/')
    if len(date_arr) > 1:
        return date_arr[2]
    return join_date

# Reset Error Log
with open("errorlog.txt", "w") as f:
    f.write("")
f.close()

#Load the .env file (must be in same directory as this file)
load_dotenv()

#Read in the data from the .env file
file_dir = os.getenv('dir')
import_addr = file_dir + '\\' + os.getenv('import_file')
skip_addr = file_dir + '\\' + os.getenv('skip_file')

# Creates a dictionary from the CSV data
with open(import_addr, newline='') as csv_file:
    import_data = []
    reader = csv.DictReader(csv_file)
    for row in reader:
        import_data.append(row)
csv_file.close()

# Gets a set of addresses to skip
skip_entries = set()
if os.path.exists(skip_addr):
    with open(skip_addr, newline='') as csv_file:
        reader = csv.reader(csv_file, delimiter=',')
        for row in reader:
            skip_entries.add(row[0])
    csv_file.close()

# Sets up the fields that will turn into a bulleted list
bullet_fields = ['Entry Feature 1', 'Entry Feature 2', 'Entry Feature 3', 'Entry Feature 4', 'Entry Feature 5', 'Entry Feature 6']
table_data = [
    {'label' : 'Finished SqFt (Not Including Garage)', 'field' : 'Living SF'},
    {'label' : 'Total SqFt (Not Including Garage)', 'field' : 'Total SF'},
    {'label' : 'Garage SqFt', 'field' : 'Garage SF'},
    {'label' : 'Number of Bedrooms', 'field' : 'Bedrooms'},
    {'label' : 'Number of Baths', 'field' : 'Bathrooms'}
]

# Writes Word documents
for entry in import_data:
    # If the Word doc field hasn't been set to 'yes' in MZ
    # and the adddress isn't in the Skip CSV file
    if entry['Entry Street Address'] not in skip_entries and entry['Word Doc Done'] == 'No':
        
        # Create a Word doc
        document = Document()
        
        # Write entry number placeholder
        paragraph = document.add_paragraph()
        paragraph.add_run('#').bold = True
        run = paragraph.add_run('xx')
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.bold = True
        
        # Write price
        paragraph = document.add_paragraph()
        paragraph.add_run(entry['Price'])
        
        # Write company info block
        paragraph = document.add_paragraph()
        paragraph.style = 'No Spacing'
        paragraph.add_run(entry['Attendee Company']).bold = True
        paragraph = document.add_paragraph()
        paragraph.style = 'No Spacing'
        paragraph.add_run(entry['Phone'])
        paragraph = document.add_paragraph()
        paragraph.add_run(entry['Website'])
        
        # Write the address
        paragraph = document.add_paragraph()
        paragraph.add_run(entry['Entry Street Address'] + ' {0} '.format(em_dash) + entry['Entry City']).bold = True
        
        # Write subdivision
        paragraph = document.add_paragraph()
        paragraph.add_run('Subdivision: ').bold = True
        paragraph.add_run(entry['Entry Subdivision'])
        
        # Write description
        paragraph = document.add_paragraph()
        paragraph.add_run('Description: ').bold = True
        paragraph.add_run(entry['Description'])
        
        # Add bulleted list of features
        for field in bullet_fields:
            if len(entry[field]) > 0:
                document.add_paragraph(entry[field], style='List Bullet')
        
        if entry['Entry Type'] == 'Single-Family Home' or entry['Entry Type'] == 'Townhome Unit':
            # Write home style
            paragraph = document.add_paragraph()
            paragraph.add_run('Home Style: ').bold = True
            paragraph.add_run(entry['Home Style'])
        
            # Write exterior
            paragraph = document.add_paragraph()
            paragraph.add_run('Exterior: ').bold = True
            paragraph.add_run(entry['Exterior'])
        
            # Write HERS
            paragraph = document.add_paragraph()
            paragraph.add_run('HERS: ').bold = True
            if len(entry['HERS Rating']) != 0:
                paragraph.add_run(entry['HERS Rating'])
            else:
                run = paragraph.add_run('Not Yet Assessed')
                run.font.color.rgb = RGBColor(255, 0, 0)
        
            # Create table
            table = document.add_table(rows=5, cols=2)
            for i, data in enumerate(table_data):
                row = table.rows[i].cells
                row[0].text = data['label']
                table_value = entry[data['field']]
                if i > 2:
                    table_value = trailing_zeros(table_value)
                row[1].text = table_value
        
            # Write contractor license number
            paragraph = document.add_paragraph()
            paragraph = document.add_paragraph()
            paragraph.add_run('Contractor License #' + entry['Contractor License']).italic = True
        
        else:
            # Write entry type
            paragraph = document.add_paragraph()
            paragraph.add_run('Entry Type: ').bold = True
            paragraph.add_run(entry['Entry Type'])
        
        # Write member since info
        paragraph = document.add_paragraph()
        join_yr = get_join_yr(entry['Member Join Date'])
        paragraph.add_run('This company has been a Member since ' + join_yr + '.').italic = True
        
        # Save the document
        document.save(file_dir + '\\' + entry['Entry Street Address'] + ' - Magazine Info.docx')